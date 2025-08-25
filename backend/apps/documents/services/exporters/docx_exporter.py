import io
import re
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from markdown_it import MarkdownIt

def export_document_to_docx(django_doc, citation_style="apa", include_comments=False):
    """Export a Django Document model instance to DOCX format with improved formatting and citation handling."""

    # ✅ ensure we got the right type of object
    from apps.documents.models import Document  # your Django model
    if not isinstance(django_doc, Document):
        raise TypeError(f"Expected Django Document model, got {type(django_doc)}")

    docx = DocxDocument()
    
    # Set up styles
    setup_document_styles(docx)
    
    # Add title
    title_para = docx.add_heading("AI Concept Note", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata if present
    # if django_doc.meta:
    #     docx.add_paragraph()  # Add space
    #     for k, v in django_doc.meta.items():
    #         p = docx.add_paragraph(f"{k}: {v}")
    #         p.runs[0].font.size = Pt(11)
    #         p.runs[0].italic = True

    # Initialize markdown parser with table support
    md = MarkdownIt("commonmark", {"linkify": False, "html": True}).enable("table")

    # Track all citations to avoid duplicates
    all_citations = set()
    
    # Process each section (must be a related manager/queryset)
    for sec in django_doc.sections.order_by("order"):
        # Add section heading
        # docx.add_heading(sec.title, level=1)
        
        content = sec.get_content() or ""
        
        # Process the content
        process_markdown_content(docx, content, md)
        
        # Collect unique citations
        for citation in sec.citations.all():
            all_citations.add(citation.reference_text)
        
        # Add comments if requested
        if include_comments:
            comments = sec.comments.filter(resolved=False)
            if comments:
                docx.add_paragraph()  # Add space
                comment_heading = docx.add_paragraph("Comments:")
                comment_heading.runs[0].bold = True
                comment_heading.runs[0].font.size = Pt(12)
                
                for cm in comments:
                    p = docx.add_paragraph(f"• {cm.author}: {cm.content}")
                    p.runs[0].italic = True
                    p.runs[0].font.size = Pt(10)

    # Add references section if there are citations
    if all_citations:
        docx.add_page_break()
        docx.add_heading("References", level=1)
        
        # Sort citations alphabetically
        sorted_citations = sorted(all_citations)
        for citation_text in sorted_citations:
            p = docx.add_paragraph(citation_text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent

    bio = io.BytesIO()
    docx.save(bio)
    bio.seek(0)
    return bio

def setup_document_styles(docx):
    """Set up document-wide styles."""
    # Set default font
    style = docx.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

def process_markdown_content(docx, content, md):
    """Process markdown content and convert to DOCX elements."""
    tokens = md.parse(content)
    
    i = 0
    while i < len(tokens):
        token = tokens[i]
        
        if token.type == 'heading_open':
            level = int(token.tag[1])  # Extract number from h1, h2, etc.
            i += 1
            if i < len(tokens) and tokens[i].type == 'inline':
                docx.add_heading(tokens[i].content, level=level + 1)  # Adjust level
            i += 1  # Skip heading_close
            
        elif token.type == 'paragraph_open':
            i += 1
            if i < len(tokens) and tokens[i].type == 'inline':
                content_text = tokens[i].content
                if content_text.strip():  # Only add non-empty paragraphs
                    para = docx.add_paragraph()
                    process_inline_text(para, content_text)
            i += 1  # Skip paragraph_close
            
        elif token.type == 'table_open':
            i = process_table(docx, tokens, i)
            
        elif token.type == 'blockquote_open':
            i = process_blockquote(docx, tokens, i)
            
        elif token.type == 'bullet_list_open':
            i = process_list(docx, tokens, i, ordered=False)
            
        elif token.type == 'ordered_list_open':
            i = process_list(docx, tokens, i, ordered=True)
            
        else:
            i += 1

def process_inline_text(paragraph, text):
    """Process inline text with bold, italic, and other formatting."""
    # First handle bold text (**text**)
    parts = re.split(r'(\*\*[^*]+?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            # Remove the ** markers and make text bold
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Now handle italic text (*text*) in remaining parts
            # Also handle escaped asterisks (\*)
            part = part.replace('\\*', '*')  # Convert escaped asterisks to regular asterisks
            italic_parts = re.split(r'(\*[^*]+?\*)', part)
            for italic_part in italic_parts:
                if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) > 2 and not italic_part.startswith('**'):
                    # Remove the * markers and make text italic
                    italic_text = italic_part[1:-1]
                    run = paragraph.add_run(italic_text)
                    run.italic = True
                else:
                    # Regular text - clean up any remaining formatting markers
                    clean_text = clean_remaining_markers(italic_part)
                    if clean_text:  # Only add non-empty text
                        paragraph.add_run(clean_text)

def clean_remaining_markers(text):
    """Clean up any remaining markdown markers in text."""
    # Remove any stray \* sequences
    text = text.replace('\\*\\*', '**').replace('\\*', '*')
    return text

def process_table(docx, tokens, start_index):
    """Process a markdown table and convert to DOCX table."""
    i = start_index + 1  # Skip table_open
    
    # Find table structure
    headers = []
    rows = []
    current_row = []
    
    while i < len(tokens) and tokens[i].type != 'table_close':
        token = tokens[i]
        
        if token.type == 'th_open':
            i += 1
            if i < len(tokens) and tokens[i].type == 'inline':
                headers.append(clean_formatting_markers(tokens[i].content))
            i += 1  # Skip th_close
            
        elif token.type == 'td_open':
            i += 1
            if i < len(tokens) and tokens[i].type == 'inline':
                current_row.append(clean_formatting_markers(tokens[i].content))
            i += 1  # Skip td_close
            
        elif token.type == 'tr_close':
            if current_row:
                rows.append(current_row)
                current_row = []
            i += 1
            
        else:
            i += 1
    
    # Create the table
    if headers:
        table = docx.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers with formatting
        header_row = table.rows[0]
        for j, header in enumerate(headers):
            cell = header_row.cells[j]
            p = cell.paragraphs[0]
            p.clear()  # Clear existing content
            # Headers are typically just bold text. This simplifies logic and avoids formatting conflicts.
            run = p.add_run(header)
            run.bold = True
        
        # Add data rows with formatting
        for row_data in rows:
            row = table.add_row()
            for j, cell_data in enumerate(row_data[:len(headers)]):  # Ensure we don't exceed columns
                cell = row.cells[j]
                cell.paragraphs[0].clear()  # Clear existing content
                process_inline_text(cell.paragraphs[0], cell_data)
    
    return i + 1  # Skip table_close

def process_blockquote(docx, tokens, start_index):
    """Process a blockquote."""
    i = start_index + 1  # Skip blockquote_open
    
    quote_text = ""
    while i < len(tokens) and tokens[i].type != 'blockquote_close':
        if tokens[i].type == 'inline':
            quote_text += tokens[i].content + " "
        i += 1
    
    if quote_text.strip():
        para = docx.add_paragraph(quote_text.strip())
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        for run in para.runs:
            run.italic = True
    
    return i + 1  # Skip blockquote_close

def process_list(docx, tokens, start_index, ordered=False):
    """Process a list (ordered or unordered)."""
    i = start_index + 1  # Skip list_open
    
    while i < len(tokens) and tokens[i].type not in ['bullet_list_close', 'ordered_list_close']:
        if tokens[i].type == 'list_item_open':
            i += 1
            if i < len(tokens) and tokens[i].type == 'paragraph_open':
                i += 1
                if i < len(tokens) and tokens[i].type == 'inline':
                    list_text = tokens[i].content
                    if ordered:
                        para = docx.add_paragraph(style='List Number')
                    else:
                        para = docx.add_paragraph(style='List Bullet')
                    process_inline_text(para, list_text)
        i += 1
    
    return i + 1  # Skip list_close

def clean_formatting_markers(text):
    """Clean up text by removing markdown formatting markers for plain text contexts."""
    # Remove bold markers
    cleaned = re.sub(r'\*\*([^*]+?)\*\*', r'\1', text)
    # Remove italic markers
    cleaned = re.sub(r'\*([^*]+?)\*', r'\1', cleaned)
    # Remove escaped asterisks
    cleaned = cleaned.replace('\\*', '*')
    return cleaned

def clean_reference_text(text):
    """Clean up reference text by removing duplicate chunks and formatting properly."""
    # Remove "Knowledge Base:" prefixes and chunk information
    cleaned = re.sub(r'Knowledge Base: [^(]+\(Chunk \d+\)\s*', '', text)
    
    # Remove extra whitespace
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    
    return cleaned if cleaned else text